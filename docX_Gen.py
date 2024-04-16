#doX_Gen.py
"""
DocX User Credentials Generator
Author: Nathaniel Jones
Version: 1.0
Description: This script creates word document for user credentials from 3 inputs
"""

# Python Imports
import tkinter as tk
from tkinter import messagebox
from docx import Document
from tkinter import PhotoImage


# Define word document function
def generate_document():
    first_name = first_name_entry.get()
    last_name = last_name_entry.get()
    password = password_entry.get()

    if not first_name or not last_name or not password:
        messagebox.showerror("Error", "Please fill in all fields.")
        return
    # Extract first initial from first_name
    first_initial = first_name[0]

    doc = Document()
    doc.add_heading(f'Welcome to [Company Name Here] {first_name} {last_name}', level=1)
    doc.add_paragraph("")  # Empty line
    doc.add_paragraph("")  # Empty line
    doc.add_paragraph(f"[Company Name Here]Email: {first_initial}.{last_name}@[compay domain].com")
    doc.add_paragraph(f"Desktop Username/Login: {first_initial}.{last_name}")
    doc.add_paragraph(f"[Company Name Here] Password: {password}")
    doc.add_paragraph(f"[Company Name Here] Username: {first_initial}.{last_name}@[compay domain]")
    doc.add_paragraph(f"[Company Name Here] Password: {password}")
    doc.add_paragraph("")  # Empty line
    doc.add_paragraph("Voicemail Password is your phone extension followed by a 0") 
    doc.add_paragraph("")  # Empty line
    doc.add_paragraph("A Technician as completed the following to your workstation:")  
    doc.add_paragraph("")  # Empty line
    doc.add_paragraph("⚬    Logged in and configured Outlook/Email")
    doc.add_paragraph("⚬    Logged in and configured MS Teams")  
    doc.add_paragraph("⚬    Logged in and Launched [Application]") 
    doc.add_paragraph("⚬    Configured [software], [software], and [software]") 
    doc.add_paragraph("⚬    Installed Default Printers")
    doc.add_paragraph("⚬    Tested Scanner(If Applicable)") 
    doc.add_paragraph("⚬    Created a desktop icon for email access") 
    doc.add_paragraph("")  # Empty line
    doc.add_paragraph("Wifi: [SSID]: [Employee password]") 

    doc.save(f"Welcome_{first_name}.docx")
    
    messagebox.showinfo("Success", "Document created successfully. **Double Check username available in AD!**")

# Define window object
def center_window(window, width, height):
    screen_width = window.winfo_screenwidth()
    screen_height = window.winfo_screenheight()
    x = (screen_width - width) // 2
    y = (screen_height - height) // 2
    return f"{width}x{height}+{x}+{y}"

# Create the main GUI window
root = tk.Tk()
root.title("DocX User Credentials Generator")

# Set window dimensions and position
window_geometry = center_window(root, 400, 600)
root.geometry(window_geometry)

# Load and display an image
image = PhotoImage(file="assets/img/ccLogoS.png")  # Replace "image.png" with your image file of your choice this will be for the logo
image = image.subsample(9)  # Change the factor to resize the image (e.g., subsample(9) divides the size)
image_label = tk.Label(root, image=image)
image_label.pack()

# Create labels and entry fields
first_name_label = tk.Label(root, text="First Name:")
first_name_label.pack()
first_name_entry = tk.Entry(root)
first_name_entry.pack()

last_name_label = tk.Label(root, text="Last Name:")
last_name_label.pack()
last_name_entry = tk.Entry(root)
last_name_entry.pack()

password_label = tk.Label(root, text="Password:")
password_label.pack()
password_entry = tk.Entry(root, show="*") # hides pw entry show="*"
password_entry.pack()

# Create the generate button
generate_button = tk.Button(root, text="Generate DocX", command=generate_document)
generate_button.pack()
# Start the GUI event loop
root.mainloop()

